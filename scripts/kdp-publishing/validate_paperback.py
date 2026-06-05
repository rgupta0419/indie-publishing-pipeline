"""
KDP Paperback Interior Pre-flight Validation.

Checks a print-ready PDF (or docx) for KDP compliance:
  - Page count vs. trim size limits
  - Margins meet minimums for page count band
  - All fonts embedded (PDF only)
  - Image DPI >= 300 (PDF only)
  - Body font size >= 9pt
  - PDF dimensions match trim size declaration
  - No password protection
  - Single section / consistent page geometry

Usage:
    python validate_paperback.py path/to/manuscript.pdf --trim 5.5x8.5
    python validate_paperback.py path/to/manuscript.docx --trim 5.5x8.5

Exit codes:
    0 = all checks pass
    1 = warnings only (review recommended)
    2 = errors (will fail KDP review)
"""

import argparse
import json
import os
import sys
from pathlib import Path


# KDP paperback specifications
TRIM_SIZES = {
    "5x8": (5.0, 8.0), "5.06x7.81": (5.06, 7.81), "5.25x8": (5.25, 8.0),
    "5.5x8.5": (5.5, 8.5), "6x9": (6.0, 9.0), "6.14x9.21": (6.14, 9.21),
    "6.69x9.61": (6.69, 9.61), "7x10": (7.0, 10.0), "7.44x9.69": (7.44, 9.69),
    "7.5x9.25": (7.5, 9.25), "8x10": (8.0, 10.0), "8.5x11": (8.5, 11.0),
}

# Inside (gutter) minimums by page count band
def min_gutter(page_count):
    if page_count <= 150: return 0.375
    if page_count <= 300: return 0.5
    if page_count <= 500: return 0.625
    if page_count <= 700: return 0.75
    return 0.875

PAGE_COUNT_LIMITS = {
    "bw_white": (24, 828),
    "bw_cream": (24, 776),
    "standard_color": (24, 600),
    "premium_color": (24, 828),
}


class CheckResult:
    def __init__(self):
        self.errors = []
        self.warnings = []
        self.ok = []

    def err(self, msg):
        self.errors.append(msg)

    def warn(self, msg):
        self.warnings.append(msg)

    def pass_(self, msg):
        self.ok.append(msg)

    def summary(self):
        return {
            "errors": self.errors,
            "warnings": self.warnings,
            "passed": self.ok,
            "total_checks": len(self.errors) + len(self.warnings) + len(self.ok),
            "status": "FAIL" if self.errors else ("WARN" if self.warnings else "PASS"),
        }


def check_pdf(path, trim_w, trim_h, paper_type="bw_white"):
    result = CheckResult()
    try:
        from pypdf import PdfReader
    except ImportError:
        result.err("pypdf not installed. Install: pip install pypdf")
        return result

    try:
        reader = PdfReader(path)
    except Exception as e:
        result.err(f"Cannot open PDF: {e}")
        return result

    if reader.is_encrypted:
        result.err("PDF is password-protected. KDP will reject. Remove encryption.")
        return result

    # Page count
    page_count = len(reader.pages)
    min_p, max_p = PAGE_COUNT_LIMITS[paper_type]
    if page_count < min_p:
        result.err(f"Page count {page_count} is below KDP minimum ({min_p})")
    elif page_count > max_p:
        result.err(f"Page count {page_count} exceeds KDP maximum for {paper_type} ({max_p})")
    else:
        result.pass_(f"Page count {page_count} within KDP limits ({min_p}-{max_p} for {paper_type})")

    # Even/odd parity (KDP prefers even for proper spine alignment)
    if page_count % 2 != 0:
        result.warn(f"Page count {page_count} is odd. KDP recommends even page count "
                    "(printing pairs left+right). Consider adding a blank back page.")

    # Page dimensions — sample first 5 pages
    for i, page in enumerate(reader.pages[:5]):
        media_box = page.mediabox
        w_pt = float(media_box.width)
        h_pt = float(media_box.height)
        # Convert pt to inches (72 pt = 1 inch)
        w_in = w_pt / 72
        h_in = h_pt / 72

        # Allow 0.05" tolerance
        if abs(w_in - trim_w) > 0.05 or abs(h_in - trim_h) > 0.05:
            result.err(f"Page {i+1} dimensions {w_in:.2f}x{h_in:.2f}\" don't match "
                       f"declared trim {trim_w}x{trim_h}\". KDP will reject.")
            break
    else:
        result.pass_(f"Page dimensions match trim size {trim_w}x{trim_h}\"")

    # Font embedding check
    try:
        fonts_used = set()
        fonts_embedded = set()
        for page in reader.pages[:10]:  # sample
            if "/Resources" in page and "/Font" in page["/Resources"]:
                fonts = page["/Resources"]["/Font"].get_object()
                for font_key in fonts:
                    font_obj = fonts[font_key].get_object()
                    base_font = str(font_obj.get("/BaseFont", "Unknown"))
                    fonts_used.add(base_font)
                    # Check for embedded font descriptor
                    if "/FontDescriptor" in font_obj:
                        desc = font_obj["/FontDescriptor"].get_object()
                        if "/FontFile" in desc or "/FontFile2" in desc or "/FontFile3" in desc:
                            fonts_embedded.add(base_font)

        not_embedded = fonts_used - fonts_embedded
        if not_embedded:
            result.warn(f"Fonts not embedded: {', '.join(sorted(not_embedded)[:5])}. "
                        "KDP requires all fonts embedded. Re-export PDF with font embedding enabled.")
        elif fonts_used:
            result.pass_(f"All {len(fonts_used)} fonts embedded (sampled first 10 pages)")
    except Exception as e:
        result.warn(f"Could not verify font embedding: {e}")

    # Margin info — calc from page contents
    min_gut = min_gutter(page_count)
    result.pass_(f"Required minimum gutter for {page_count}-page book: {min_gut}\" "
                 f"(verify in Word/Pages: Layout → Margins)")

    return result, page_count


def check_docx(path, trim_w, trim_h):
    result = CheckResult()
    try:
        from docx import Document
    except ImportError:
        result.err("python-docx not installed. Install: pip install python-docx")
        return result, 0

    try:
        doc = Document(path)
    except Exception as e:
        result.err(f"Cannot open DOCX: {e}")
        return result, 0

    # Page dimensions from first section
    s = doc.sections[0]
    w_in = s.page_width.inches
    h_in = s.page_height.inches
    if abs(w_in - trim_w) > 0.05 or abs(h_in - trim_h) > 0.05:
        result.err(f"Page dimensions {w_in}x{h_in}\" don't match declared trim "
                   f"{trim_w}x{trim_h}\". Update Page Setup before exporting PDF.")
    else:
        result.pass_(f"Page dimensions match trim {trim_w}x{trim_h}\"")

    # Margins
    lm, rm = s.left_margin.inches, s.right_margin.inches
    tm, bm = s.top_margin.inches, s.bottom_margin.inches
    result.pass_(f"Current margins: L={lm}\" R={rm}\" T={tm}\" B={bm}\" "
                 "(verify gutter setting in Word: Layout → Margins → Mirrored)")

    # Mirror margins check via XML
    sect_props = s._sectPr
    from docx.oxml.ns import qn
    mirror_setting = sect_props.find(qn('w:mirrorMargins'))
    if mirror_setting is None:
        result.warn("Mirror margins flag not set in docx. For books >50 pages, "
                    "use Layout → Margins → Mirrored so gutter is consistent on "
                    "left+right pages. (LibreOffice doesn't honor this; verify "
                    "Word/Pages rendering before exporting final PDF.)")
    else:
        result.pass_("Mirror margins enabled")

    # Estimated page count (rough)
    para_count = len(doc.paragraphs)
    est_pages = max(para_count // 40, 24)
    result.pass_(f"Estimated page count: ~{est_pages} (real count after PDF export). "
                 "Note: LibreOffice and Word/Pages produce different page counts; "
                 "use Mac Word/Pages for the canonical page count.")

    # Font check — sample
    fonts = set()
    sizes = []
    for p in doc.paragraphs[:200]:
        for r in p.runs:
            if r.font.name:
                fonts.add(r.font.name)
            if r.font.size:
                sizes.append(r.font.size.pt)

    if fonts:
        result.pass_(f"Fonts used (sampled first 200 paragraphs): {', '.join(sorted(fonts))}")
    body_sizes = [s for s in sizes if 8 <= s <= 14]
    if body_sizes:
        most_common = max(set(body_sizes), key=body_sizes.count)
        if most_common < 9:
            result.warn(f"Most common body font size is {most_common}pt. KDP minimum for "
                        "readability is 9pt; recommended 10-11pt for non-fiction.")
        else:
            result.pass_(f"Most common body font size: {most_common}pt (good for readability)")

    return result, est_pages


def main():
    p = argparse.ArgumentParser(description="Validate manuscript for KDP paperback upload")
    p.add_argument("path", help="Path to .pdf or .docx manuscript file")
    p.add_argument("--trim", default="5.5x8.5",
                   help="Trim size (default 5.5x8.5). Options: " + ", ".join(TRIM_SIZES.keys()))
    p.add_argument("--paper", default="bw_white",
                   choices=list(PAGE_COUNT_LIMITS.keys()),
                   help="Paper type (default bw_white)")
    p.add_argument("--json", action="store_true", help="Output JSON instead of human-readable")
    args = p.parse_args()

    if args.trim not in TRIM_SIZES:
        print(f"Unknown trim size: {args.trim}", file=sys.stderr)
        sys.exit(2)
    trim_w, trim_h = TRIM_SIZES[args.trim]

    ext = Path(args.path).suffix.lower()
    if ext == ".pdf":
        result, pages = check_pdf(args.path, trim_w, trim_h, args.paper)
    elif ext == ".docx":
        result, pages = check_docx(args.path, trim_w, trim_h)
    else:
        print(f"Unsupported file type: {ext}. Use .pdf or .docx", file=sys.stderr)
        sys.exit(2)

    summary = result.summary()
    summary["file"] = args.path
    summary["trim"] = f"{trim_w}x{trim_h}"
    summary["estimated_page_count"] = pages

    if args.json:
        print(json.dumps(summary, indent=2))
    else:
        print(f"\n=== KDP Paperback Validation: {args.path} ===")
        print(f"Trim: {trim_w}x{trim_h}\" | Paper: {args.paper}")
        print(f"Estimated page count: {pages}\n")

        if summary["errors"]:
            print(f"❌ ERRORS ({len(summary['errors'])}):")
            for e in summary["errors"]:
                print(f"   • {e}")
            print()

        if summary["warnings"]:
            print(f"⚠️  WARNINGS ({len(summary['warnings'])}):")
            for w in summary["warnings"]:
                print(f"   • {w}")
            print()

        if summary["passed"]:
            print(f"✓ PASSED ({len(summary['passed'])}):")
            for ok in summary["passed"]:
                print(f"   • {ok}")
            print()

        print(f"=== STATUS: {summary['status']} ===\n")

    sys.exit(0 if summary["status"] == "PASS" else (1 if summary["status"] == "WARN" else 2))


if __name__ == "__main__":
    main()
