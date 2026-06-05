"""
Stage 5 — Apply surgical fixes for CX issues identified in the audit.

Three fix types:
1. Heading orphans: keep_with_next on heading + keep_lines on heading and next 2 body paragraphs
2. Callout-box splits: cantSplit on specific tables (NOT all — that adds 10+ pages)
3. Author closing-block orphans: keep_with_next chain on the closing paragraphs

Apply selectively — pass paragraph indices and table indices for targeted fixes.

For callout-box splits where the box is too big to fit one page, shorten the bullet text
manually (open the docx, find the bullet, edit). cantSplit only helps for small-to-medium boxes.
"""

import argparse
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def apply_keep_with_next(para):
    pPr = para._element.get_or_add_pPr()
    if pPr.find(qn('w:keepNext')) is None:
        pPr.append(OxmlElement('w:keepNext'))


def apply_keep_lines(para):
    pPr = para._element.get_or_add_pPr()
    if pPr.find(qn('w:keepLines')) is None:
        pPr.append(OxmlElement('w:keepLines'))


def apply_widow_control(para):
    pPr = para._element.get_or_add_pPr()
    if pPr.find(qn('w:widowControl')) is None:
        pPr.append(OxmlElement('w:widowControl'))


def apply_cantSplit_to_table(tbl):
    for row in tbl.rows:
        trPr = row._tr.find(qn('w:trPr'))
        if trPr is None:
            trPr = OxmlElement('w:trPr')
            row._tr.insert(0, trPr)
        if trPr.find(qn('w:cantSplit')) is None:
            trPr.append(OxmlElement('w:cantSplit'))


def fix_heading_orphans(doc, heading_indices):
    """For each paragraph index that's an orphaned heading, apply keep_with_next + keep_lines.
    Also apply keep_lines to the next 2 body paragraphs after the heading."""
    for hi in heading_indices:
        if hi >= len(doc.paragraphs):
            continue
        p = doc.paragraphs[hi]
        apply_keep_with_next(p)
        apply_keep_lines(p)
        # Next 2 body paragraphs
        for offset in (1, 2):
            if hi + offset < len(doc.paragraphs):
                np = doc.paragraphs[hi + offset]
                if np.text.strip():
                    apply_keep_lines(np)


def fix_callout_splits(doc, table_indices):
    """For each table index, apply cantSplit on all rows. Only do this for tables small
    enough to fit on a single page — otherwise it pushes empty space."""
    for ti in table_indices:
        if ti >= len(doc.tables):
            continue
        apply_cantSplit_to_table(doc.tables[ti])


def fix_widow_control_global(doc):
    """Apply widowControl to every paragraph in the document. Helps prevent 1-line widows
    in regular body text."""
    def visit(paragraphs):
        for p in paragraphs:
            apply_widow_control(p)
    visit(doc.paragraphs)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                visit(cell.paragraphs)


def fix_table_preceding_paragraphs(doc):
    """Apply keep_with_next to any paragraph immediately preceding a table.
    This keeps the section heading attached to its callout box."""
    body = doc._body._element
    prev_para = None
    for child in body:
        tag = child.tag.split('}')[-1]
        if tag == 'p':
            prev_para = child
        elif tag == 'tbl' and prev_para is not None:
            pPr = prev_para.find(qn('w:pPr'))
            if pPr is None:
                pPr = OxmlElement('w:pPr')
                prev_para.insert(0, pPr)
            if pPr.find(qn('w:keepNext')) is None:
                pPr.append(OxmlElement('w:keepNext'))
            prev_para = None


def fix_signature_block(doc, marker_text="", chain_size=3):
    """Find a closing-block marker paragraph and chain keep_with_next through next N
    non-empty paragraphs. Useful for keeping author signature with closing text."""
    found_at = None
    for i, p in enumerate(doc.paragraphs):
        if marker_text in p.text:
            found_at = i
            break
    if found_at is None:
        return
    chain = [doc.paragraphs[found_at]]
    j = found_at + 1
    while len(chain) < chain_size and j < len(doc.paragraphs):
        if doc.paragraphs[j].text.strip():
            chain.append(doc.paragraphs[j])
        j += 1
    for p in chain[:-1]:  # all except last get keep_with_next
        apply_keep_with_next(p)
    for p in chain:
        apply_keep_lines(p)


def main():
    p = argparse.ArgumentParser()
    p.add_argument("source")
    p.add_argument("output")
    p.add_argument("--heading-orphans", default="", help="Comma-separated paragraph indices that are orphaned headings")
    p.add_argument("--callout-splits", default="", help="Comma-separated table indices to apply cantSplit to")
    p.add_argument("--global-widow", action="store_true", help="Apply widowControl to all paragraphs")
    p.add_argument("--keep-tables-with-headings", action="store_true", help="Apply keep_with_next on all paragraphs preceding tables")
    p.add_argument("--signature-marker", default="", help="Text marker for the signature closing block (e.g., 'Welcome.', 'With gratitude,')")
    args = p.parse_args()

    doc = Document(args.source)

    if args.heading_orphans:
        idx = [int(x.strip()) for x in args.heading_orphans.split(',') if x.strip()]
        fix_heading_orphans(doc, idx)
        print(f"Applied keep_with_next + keep_lines to {len(idx)} heading paragraphs")

    if args.callout_splits:
        idx = [int(x.strip()) for x in args.callout_splits.split(',') if x.strip()]
        fix_callout_splits(doc, idx)
        print(f"Applied cantSplit to {len(idx)} tables")

    if args.global_widow:
        fix_widow_control_global(doc)
        print("Applied widowControl globally")

    if args.keep_tables_with_headings:
        fix_table_preceding_paragraphs(doc)
        print("Applied keep_with_next to all paragraphs preceding tables")

    if args.signature_marker:
        fix_signature_block(doc, args.signature_marker)
        print(f"Fixed signature block starting at: '{args.signature_marker}'")

    doc.save(args.output)
    print(f"Saved: {args.output}")


if __name__ == "__main__":
    main()
