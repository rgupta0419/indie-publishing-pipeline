"""
Stage 3 — Reflow a master manuscript from US Letter / Georgia → trade trim / DejaVu Serif.

Usage:
    python reflow_to_trim.py <source.docx> <output.docx> [--trim 5.5x8.5] [--font "DejaVu Serif"]

What this does:
- Resets page size and margins to target trim (default 5.5×8.5")
- Sets mirror margins (binding gutter on inner side)
- Swaps body font globally (preserves bold/italic/color)
- Scales body text from 11.5pt → 11pt proportionally
- Resizes inline images to fit new text width
- Resizes all tables (callout boxes) to fit page width

After running, convert via LibreOffice:
    libreoffice --headless --convert-to pdf <output.docx>
    gs -o opt.pdf -sDEVICE=pdfwrite -dPDFSETTINGS=/printer -dEmbedAllFonts=true -dSubsetFonts=true <output>.pdf
"""

import argparse
from docx import Document
from docx.shared import Inches, Pt, Emu
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def parse_trim(s):
    """Parse '5.5x8.5' → (5.5, 8.5)."""
    w, h = s.lower().split('x')
    return float(w), float(h)


def reflow(source_path, output_path, trim_w=5.5, trim_h=8.5,
           target_font="DejaVu Serif",
           margin_top=0.75, margin_bottom=0.75,
           margin_inner=0.875, margin_outer=0.625):
    """Reflow a docx to trade-trim.

    Margins use the standard mirror-margin convention: inner (gutter) wider than outer.
    For 5.5×8.5 with 250 pages, defaults are right; for thicker books increase inner to 1.0".
    """
    doc = Document(source_path)

    # 1. Reset all sections to target trim with mirror margins
    for sec in doc.sections:
        sec.page_width = Inches(trim_w)
        sec.page_height = Inches(trim_h)
        sec.top_margin = Inches(margin_top)
        sec.bottom_margin = Inches(margin_bottom)
        sec.left_margin = Inches(margin_inner)
        sec.right_margin = Inches(margin_outer)
        sec.gutter = Inches(0.0)
        # Enable mirror margins
        sectPr = sec._sectPr
        if sectPr.find(qn('w:mirrorMargins')) is None:
            mm = OxmlElement('w:mirrorMargins')
            sectPr.append(mm)

    # 2. Helper — set font on a run via XML
    def set_font_xml(run, font_name):
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.insert(0, rFonts)
        for attr in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
            rFonts.set(qn(attr), font_name)

    # 3. Walk all paragraphs (body + tables + headers/footers), swap font and rescale sizes
    runs_changed = 0
    sizes_changed = 0

    def scale_size(orig_pt):
        if orig_pt is None:
            return None
        # Body sizes (10.5–12.5) → 11pt
        if 10.5 <= orig_pt <= 12.5:
            return 11.0
        # Headings — keep proportional reduction ~5%
        return round(orig_pt * 0.95, 1)

    def process_run(run):
        nonlocal runs_changed, sizes_changed
        cur_font = run.font.name
        if cur_font is None or "Georgia" in (cur_font or "") or "Charis" in (cur_font or ""):
            set_font_xml(run, target_font)
            runs_changed += 1
        cur_size = run.font.size
        if cur_size is not None:
            new_pt = scale_size(cur_size.pt)
            if new_pt and abs(new_pt - cur_size.pt) > 0.05:
                run.font.size = Pt(new_pt)
                sizes_changed += 1

    def walk(paragraphs):
        for para in paragraphs:
            for run in para.runs:
                process_run(run)

    walk(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                walk(cell.paragraphs)
    for section in doc.sections:
        for hdr in (section.header, section.first_page_header, section.even_page_header):
            if hdr:
                walk(hdr.paragraphs)
        for ftr in (section.footer, section.first_page_footer, section.even_page_footer):
            if ftr:
                walk(ftr.paragraphs)

    # 4. Resize inline images to fit new text width
    # Text width = trim_w - margin_inner - margin_outer
    text_width = Inches(trim_w - margin_inner - margin_outer)
    imgs_resized = 0
    for shape in doc.inline_shapes:
        if shape.width > text_width:
            ratio = float(text_width) / float(shape.width)
            new_h = Emu(int(float(shape.height) * ratio))
            shape.width = text_width
            shape.height = new_h
            imgs_resized += 1

    # 5. Resize all tables to fit text width (callout boxes were sized for old trim)
    table_max_width = Inches(trim_w - margin_inner - margin_outer - 0.2)  # 0.2" margin of safety
    for table in doc.tables:
        tblPr = table._element.find(qn('w:tblPr'))
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            table._element.insert(0, tblPr)
        # Remove any existing tblW
        for el in tblPr.findall(qn('w:tblW')):
            tblPr.remove(el)
        tblW = OxmlElement('w:tblW')
        tblW.set(qn('w:w'), str(int(table_max_width.twips)))
        tblW.set(qn('w:type'), 'dxa')
        tblPr.insert(0, tblW)
        # Equalize columns
        n_cols = len(table.columns)
        tblGrid = table._element.find(qn('w:tblGrid'))
        if tblGrid is not None and n_cols > 0:
            for gc in tblGrid.findall(qn('w:gridCol')):
                tblGrid.remove(gc)
            col_w = int(table_max_width.twips / n_cols)
            for _ in range(n_cols):
                gc = OxmlElement('w:gridCol')
                gc.set(qn('w:w'), str(col_w))
                tblGrid.append(gc)

    doc.save(output_path)
    print(f"Reflowed: {source_path} → {output_path}")
    print(f"  Trim: {trim_w}×{trim_h}\"")
    print(f"  Font swapped on {runs_changed} runs (target: {target_font})")
    print(f"  Sizes rescaled: {sizes_changed}")
    print(f"  Images resized to fit: {imgs_resized}")
    print(f"  Tables resized: {len(doc.tables)}")


if __name__ == "__main__":
    p = argparse.ArgumentParser()
    p.add_argument("source", help="Source .docx")
    p.add_argument("output", help="Output .docx (will be overwritten)")
    p.add_argument("--trim", default="5.5x8.5", help="Trim size, e.g. 5.5x8.5 or 6x9")
    p.add_argument("--font", default="DejaVu Serif", help="Target body font")
    args = p.parse_args()
    w, h = parse_trim(args.trim)
    reflow(args.source, args.output, trim_w=w, trim_h=h, target_font=args.font)
