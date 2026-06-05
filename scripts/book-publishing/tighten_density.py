"""
Stage 3 — Tighten body density toward a target words-per-page benchmark.

Use after reflow if page count is significantly higher than target (e.g., AoY-style ~360 wpp).

Default behavior:
- Body 11pt → 10.5pt
- Line spacing 1.4 → 1.25
- Outer margin 0.625" → 0.55"
- Top/bottom margin 0.75" → 0.625"
- Keeps inner margin (gutter) for binding

Result: ~20–25% fewer pages, content denser but still readable.
"""

import argparse
from docx import Document
from docx.shared import Pt, Inches


def tighten(source_path, output_path,
            new_body_size=10.5, new_line_spacing=1.25,
            margin_top=0.625, margin_bottom=0.625,
            margin_inner=0.75, margin_outer=0.55):
    doc = Document(source_path)

    # Update margins on all sections
    for sec in doc.sections:
        sec.top_margin = Inches(margin_top)
        sec.bottom_margin = Inches(margin_bottom)
        sec.left_margin = Inches(margin_inner)
        sec.right_margin = Inches(margin_outer)

    def scale_size(orig_pt):
        if orig_pt is None:
            return None
        if 10.5 <= orig_pt <= 12.5:
            return new_body_size
        return round(orig_pt * 0.95, 1)

    runs_resized = 0
    paras_relined = 0

    def process_paragraphs(paragraphs):
        nonlocal runs_resized, paras_relined
        for para in paragraphs:
            pf = para.paragraph_format
            # Tighten line spacing on body paragraphs
            if pf.line_spacing and pf.line_spacing > 1.30:
                pf.line_spacing = new_line_spacing
                paras_relined += 1
            # Rescale font sizes
            for run in para.runs:
                cs = run.font.size
                if cs is not None:
                    new_pt = scale_size(cs.pt)
                    if new_pt and abs(new_pt - cs.pt) > 0.05:
                        run.font.size = Pt(new_pt)
                        runs_resized += 1

    process_paragraphs(doc.paragraphs)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                process_paragraphs(cell.paragraphs)

    doc.save(output_path)
    print(f"Tightened: {source_path} → {output_path}")
    print(f"  Body size: {new_body_size}pt | Line spacing: {new_line_spacing}")
    print(f"  Margins: T{margin_top} B{margin_bottom} I{margin_inner} O{margin_outer}")
    print(f"  Runs resized: {runs_resized} | Paragraphs re-leaded: {paras_relined}")


if __name__ == "__main__":
    p = argparse.ArgumentParser()
    p.add_argument("source")
    p.add_argument("output")
    p.add_argument("--body-size", type=float, default=10.5)
    p.add_argument("--line-spacing", type=float, default=1.25)
    args = p.parse_args()
    tighten(args.source, args.output,
            new_body_size=args.body_size,
            new_line_spacing=args.line_spacing)
