"""
Stage 4 — Add proper book footprint: running headers, page numbers, mirror layout.

Without this, page numbers float at unpredictable positions and pages have no top anchor.
After this, every page has a consistent top (running header) and bottom (page number).

Configurable:
- Header text (default: book title in tracked small caps style)
- Header/footer color (default: light grey)
- Header/footer font sizes
- First-page suppression (title page never shows running header)
"""

import argparse
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def add_page_number_field(paragraph):
    """Insert a live PAGE field code into the paragraph."""
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    return run


def set_run_style(run, size=10, font="DejaVu Serif",
                  color=RGBColor(0x66, 0x66, 0x66), italic=False, bold=False):
    run.font.name = font
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.italic = italic
    run.font.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    for attr in ('w:ascii', 'w:hAnsi', 'w:cs'):
        rFonts.set(qn(attr), font)


def add_footprint(source_path, output_path, header_text="Y A T U",
                  header_size=9, page_num_size=10, font="DejaVu Serif",
                  header_color=(0x88, 0x88, 0x88), page_num_color=(0x66, 0x66, 0x66),
                  header_distance=0.4, footer_distance=0.4):
    doc = Document(source_path)

    for sec in doc.sections:
        sec.header_distance = Inches(header_distance)
        sec.footer_distance = Inches(footer_distance)
        # Mirror margins (in case not already set)
        sectPr = sec._sectPr
        if sectPr.find(qn('w:mirrorMargins')) is None:
            mm = OxmlElement('w:mirrorMargins')
            sectPr.append(mm)
        # Different first page (title page no header)
        sec.different_first_page_header_footer = True

    sec = doc.sections[0]
    h_color = RGBColor(*header_color)
    pn_color = RGBColor(*page_num_color)

    # Set header on all body pages — both odd ('header') and even ('even_page_header')
    for hdr_attr in ('header', 'even_page_header'):
        hdr = getattr(sec, hdr_attr)
        for p in hdr.paragraphs:
            p.clear()
        hp = hdr.paragraphs[0] if hdr.paragraphs else hdr.add_paragraph()
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hp.paragraph_format.space_before = Pt(0)
        hp.paragraph_format.space_after = Pt(0)
        run = hp.add_run(header_text)
        set_run_style(run, size=header_size, font=font, color=h_color)

    # Set footer with centered page number
    for ftr_attr in ('footer', 'even_page_footer'):
        ftr = getattr(sec, ftr_attr)
        for p in ftr.paragraphs:
            p.clear()
        fp = ftr.paragraphs[0] if ftr.paragraphs else ftr.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp.paragraph_format.space_before = Pt(0)
        fp.paragraph_format.space_after = Pt(0)
        pn_run = add_page_number_field(fp)
        set_run_style(pn_run, size=page_num_size, font=font, color=pn_color)

    # First-page header/footer: blank
    for attr in ('first_page_header', 'first_page_footer'):
        target = getattr(sec, attr)
        for p in target.paragraphs:
            p.clear()
        if not target.paragraphs:
            target.add_paragraph()

    # Enable evenAndOddHeaders in document settings
    settings = doc.settings.element
    if settings.find(qn('w:evenAndOddHeaders')) is None:
        el = OxmlElement('w:evenAndOddHeaders')
        settings.append(el)

    doc.save(output_path)
    print(f"Footprint added: {source_path} → {output_path}")
    print(f"  Header: '{header_text}' centered, {header_size}pt {font}")
    print(f"  Footer: centered page number, {page_num_size}pt")
    print(f"  Header/footer distance: {header_distance}\"")
    print(f"  First page suppressed: yes")


if __name__ == "__main__":
    p = argparse.ArgumentParser()
    p.add_argument("source")
    p.add_argument("output")
    p.add_argument("--header", default="Y A T U", help="Running header text")
    p.add_argument("--font", default="DejaVu Serif")
    args = p.parse_args()
    add_footprint(args.source, args.output, header_text=args.header, font=args.font)
