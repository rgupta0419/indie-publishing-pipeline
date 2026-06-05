"""
Stage 5b — Pre-press remediation.

Apply the hard-issue fixes from the third-party pre-press review checklist:
1. Exact trim size (force inch-based dimensions, not mm)
2. Color → B&W conversion (saves ~$1.50/copy in print cost)
3. Mirror margins (inside 0.75" / outside 0.5" / top-bottom 0.5")
4. Smart quote conversion (' " → curly equivalents)
5. Auto-hyphenation enabled (eliminates word-spacing rivers)
6. PDF metadata title fixed (was likely "Document1" or "[filename]")

Run AFTER reflow + footprint + audit, BEFORE final PDF export.

Usage:
    python pre_press_fix.py <input.docx> <output.docx> \\
        --book-title "Your Book Title" \\
        --author "Author Name" \\
        --subject "Your Book Subject Line"

After running, convert to PDF and force exact page size:
    libreoffice --headless --convert-to pdf <output>.docx
    gs -o exact.pdf -sDEVICE=pdfwrite -dDEVICEWIDTHPOINTS=396 -dDEVICEHEIGHTPOINTS=612 \\
       -dFIXEDMEDIA -dEmbedAllFonts=true -dSubsetFonts=true <output>.pdf
"""

import argparse
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def fix_quotes_in_text(text):
    """Replace straight quotes with curly ones based on context."""
    if not text:
        return text, 0
    n = 0
    out = []
    for i, ch in enumerate(text):
        if ch == "'":
            prev = text[i - 1] if i > 0 else ' '
            nxt = text[i + 1] if i + 1 < len(text) else ' '
            if prev.isalpha() and (nxt.isalpha() or not nxt.strip()):
                out.append('\u2019')
            elif not prev.strip() or prev in '({[':
                out.append('\u2018')
            else:
                out.append('\u2019')
            n += 1
        elif ch == '"':
            prev = text[i - 1] if i > 0 else ' '
            if not prev.strip() or prev in '({[':
                out.append('\u201C')
            else:
                out.append('\u201D')
            n += 1
        else:
            out.append(ch)
    return ''.join(out), n


def remediate(source_path, output_path,
              trim_w=5.5, trim_h=8.5,
              margin_top=0.5, margin_bottom=0.5,
              margin_inner=0.75, margin_outer=0.5,
              book_title=None, author=None, subject=None,
              convert_colors=True, fix_quotes=True, enable_hyphenation=True):
    doc = Document(source_path)

    # 1. EXACT TRIM (inch-based to avoid mm rounding)
    print("=== 1. Exact trim ===")
    for sec in doc.sections:
        sec.page_width = Inches(trim_w)
        sec.page_height = Inches(trim_h)
    print(f"  Set to exactly {trim_w} × {trim_h} inches on {len(doc.sections)} section(s)")

    # 2. MIRROR MARGINS
    print("\n=== 2. Mirror margins ===")
    for sec in doc.sections:
        sec.top_margin = Inches(margin_top)
        sec.bottom_margin = Inches(margin_bottom)
        sec.left_margin = Inches(margin_inner)
        sec.right_margin = Inches(margin_outer)
        sec.gutter = Inches(0)
        sectPr = sec._sectPr
        if sectPr.find(qn('w:mirrorMargins')) is None:
            sectPr.append(OxmlElement('w:mirrorMargins'))
    print(f"  Inner {margin_inner}\", Outer {margin_outer}\", Top {margin_top}\", Bottom {margin_bottom}\"")

    # 3. COLOR REMOVAL — comprehensive: text, shading, BORDERS, paragraph borders, highlights
    if convert_colors:
        print("\n=== 3. Color → B&W (comprehensive — text + shading + borders + highlights) ===")
        runs_recolored = 0

        def visit_run(run):
            nonlocal runs_recolored
            color = run.font.color
            if color and color.rgb is not None:
                rgb_int = int(str(color.rgb), 16)
                r, g, b = (rgb_int >> 16) & 0xFF, (rgb_int >> 8) & 0xFF, rgb_int & 0xFF
                max_diff = max(abs(r - g), abs(g - b), abs(r - b))
                if max_diff > 5:
                    if g > r and g > b:
                        run.font.color.rgb = RGBColor(0x70, 0x70, 0x70)
                    elif b > r and b > g:
                        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
                    else:
                        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                    runs_recolored += 1

        def walk(paragraphs):
            for p in paragraphs:
                for r in p.runs:
                    visit_run(r)

        walk(doc.paragraphs)
        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    walk(cell.paragraphs)
        for sec in doc.sections:
            for hdr in (sec.header, sec.first_page_header, sec.even_page_header):
                if hdr:
                    walk(hdr.paragraphs)
            for ftr in (sec.footer, sec.first_page_footer, sec.even_page_footer):
                if ftr:
                    walk(ftr.paragraphs)
        print(f"  Color text runs → B&W: {runs_recolored}")

        # Cell shading
        cell_shading = 0
        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    tcPr = cell._tc.find(qn('w:tcPr'))
                    if tcPr is not None:
                        shd = tcPr.find(qn('w:shd'))
                        if shd is not None:
                            fill = shd.get(qn('w:fill'))
                            if fill and fill not in ('auto', 'FFFFFF'):
                                shd.set(qn('w:fill'), 'auto')
                                cell_shading += 1
        print(f"  Cell shading (table backgrounds) removed: {cell_shading}")

        # Table borders — CRITICAL: a single colored border keeps entire book in color-print pricing
        tbl_borders = 0
        for tbl in doc.tables:
            tblPr = tbl._element.find(qn('w:tblPr'))
            if tblPr is not None:
                tblBorders = tblPr.find(qn('w:tblBorders'))
                if tblBorders is not None:
                    for border in tblBorders:
                        c = border.get(qn('w:color'))
                        if c and c not in ('auto', '000000'):
                            border.set(qn('w:color'), '000000')
                            tbl_borders += 1
            for row in tbl.rows:
                for cell in row.cells:
                    tcPr = cell._tc.find(qn('w:tcPr'))
                    if tcPr is not None:
                        tcBorders = tcPr.find(qn('w:tcBorders'))
                        if tcBorders is not None:
                            for border in tcBorders:
                                c = border.get(qn('w:color'))
                                if c and c not in ('auto', '000000'):
                                    border.set(qn('w:color'), '000000')
                                    tbl_borders += 1
        print(f"  Table border colors → black: {tbl_borders}")

        # Paragraph borders (blockquote rules, decorative lines)
        para_borders = 0
        def fix_para_borders(p):
            nonlocal para_borders
            pPr = p._element.find(qn('w:pPr'))
            if pPr is not None:
                pBdr = pPr.find(qn('w:pBdr'))
                if pBdr is not None:
                    for border in pBdr:
                        c = border.get(qn('w:color'))
                        if c and c not in ('auto', '000000'):
                            border.set(qn('w:color'), '000000')
                            para_borders += 1

        for p in doc.paragraphs:
            fix_para_borders(p)
        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        fix_para_borders(p)
        print(f"  Paragraph border colors → black: {para_borders}")

        # Run-level shading (highlight backgrounds) and highlight property
        run_shading = 0
        run_highlights = 0
        def fix_run_color_extras(run):
            nonlocal run_shading, run_highlights
            rPr = run._element.find(qn('w:rPr'))
            if rPr is not None:
                shd = rPr.find(qn('w:shd'))
                if shd is not None:
                    fill = shd.get(qn('w:fill'))
                    if fill and fill not in ('auto', 'FFFFFF'):
                        shd.set(qn('w:fill'), 'auto')
                        run_shading += 1
                hl = rPr.find(qn('w:highlight'))
                if hl is not None:
                    val = hl.get(qn('w:val'))
                    if val and val not in ('none', 'white'):
                        hl.set(qn('w:val'), 'none')
                        run_highlights += 1

        for p in doc.paragraphs:
            for r in p.runs:
                fix_run_color_extras(r)
        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for r in p.runs:
                            fix_run_color_extras(r)
        print(f"  Run shading removed: {run_shading}")
        print(f"  Run highlights removed: {run_highlights}")

    # 4. SMART QUOTES
    if fix_quotes:
        print("\n=== 4. Smart quotes ===")
        n_swaps = 0

        def fix_run(run):
            nonlocal n_swaps
            if "'" in run.text or '"' in run.text:
                new_text, n = fix_quotes_in_text(run.text)
                run.text = new_text
                n_swaps += n

        def walk_q(paragraphs):
            for p in paragraphs:
                for r in p.runs:
                    fix_run(r)

        walk_q(doc.paragraphs)
        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    walk_q(cell.paragraphs)
        print(f"  Straight quotes converted to curly: {n_swaps}")

    # 5. AUTO-HYPHENATION
    if enable_hyphenation:
        print("\n=== 5. Auto-hyphenation ===")
        settings = doc.settings.element
        if settings.find(qn('w:autoHyphenation')) is None:
            settings.append(OxmlElement('w:autoHyphenation'))
            print("  Enabled")
        else:
            print("  Already enabled")

    # 6. METADATA
    print("\n=== 6. Document properties ===")
    if book_title:
        old = doc.core_properties.title
        doc.core_properties.title = book_title
        print(f"  Title: {old!r} → {book_title!r}")
    if author:
        doc.core_properties.author = author
        print(f"  Author: {author}")
    if subject:
        doc.core_properties.subject = subject
        print(f"  Subject: {subject}")

    doc.save(output_path)
    print(f"\nSaved: {output_path}")


if __name__ == "__main__":
    p = argparse.ArgumentParser()
    p.add_argument("source")
    p.add_argument("output")
    p.add_argument("--trim", default="5.5x8.5", help="Trim, e.g. 5.5x8.5 or 6x9")
    p.add_argument("--margin-inner", type=float, default=0.75)
    p.add_argument("--margin-outer", type=float, default=0.5)
    p.add_argument("--margin-top", type=float, default=0.5)
    p.add_argument("--margin-bottom", type=float, default=0.5)
    p.add_argument("--book-title", default=None)
    p.add_argument("--author", default=None)
    p.add_argument("--subject", default=None)
    p.add_argument("--no-color-conversion", action="store_true")
    p.add_argument("--no-quote-fix", action="store_true")
    p.add_argument("--no-hyphenation", action="store_true")
    args = p.parse_args()

    w, h = [float(x) for x in args.trim.lower().split('x')]
    remediate(args.source, args.output,
              trim_w=w, trim_h=h,
              margin_top=args.margin_top, margin_bottom=args.margin_bottom,
              margin_inner=args.margin_inner, margin_outer=args.margin_outer,
              book_title=args.book_title, author=args.author, subject=args.subject,
              convert_colors=not args.no_color_conversion,
              fix_quotes=not args.no_quote_fix,
              enable_hyphenation=not args.no_hyphenation)
